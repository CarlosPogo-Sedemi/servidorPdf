import base64
from io import BytesIO
from typing import Dict, Any, List
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from PIL import Image
import re
from docx import Document
from fastapi.staticfiles import StaticFiles
#from typing import List

app = FastAPI(title="Servidor Universal de Reportes - SEDEMI")

app.mount("/assets", StaticFiles(directory="assets"), name="assets") 

# ==========================================
# MODELO DE ENTRADA GENÉRICO
# ==========================================
class PayloadUniversal(BaseModel):
    template_name: str         
    data: Dict[str, Any]


# ==========================================
# MODELOS PARA COMPRESIÓN DE FOTOS
# ==========================================
class FotoItem(BaseModel):
    NombreArchivo: str
    NovedadFoto: str
    SubirNube: bool

class LoteFotos(BaseModel):
    fotos: List[FotoItem]


# ==========================================
# ENDPOINT DE COMPRESIÓN
# ==========================================
@app.post("/api/v1/comprimir-fotos/")
def comprimir_fotos(payload: LoteFotos):
    fotos_optimizadas = []
    
    for foto in payload.fotos:
        # Solo procesamos si realmente trae foto y hay que subirla
        if foto.SubirNube and foto.NovedadFoto:
            try:
                # 1. Separar la cabecera (data:image/jpeg;base64,) del contenido
                if "base64," in foto.NovedadFoto:
                    b64_data = foto.NovedadFoto.split("base64,")[-1]
                else:
                    b64_data = foto.NovedadFoto
                    
                img_bytes = base64.b64decode(b64_data)
                img = Image.open(BytesIO(img_bytes)).convert("RGB")
                
                # 2. Redimensionar si es gigante (ancho máximo 800px)
                max_width = 800
                if img.width > max_width:
                    wpercent = (max_width / float(img.width))
                    hsize = int((float(img.height) * float(wpercent)))
                    img = img.resize((max_width, hsize), Image.Resampling.LANCZOS)
                
                # 3. Comprimir la calidad
                buffer_optimizado = BytesIO()
                # Guardamos como JPEG con calidad al 60% (Suficiente para reportes)
                img.save(buffer_optimizado, format="JPEG", optimize=True, quality=60)
                
                # 4. Reconstruir el Base64
                nueva_b64 = f"data:image/jpeg;base64,{base64.b64encode(buffer_optimizado.getvalue()).decode()}"
                foto.NovedadFoto = nueva_b64
                
            except Exception as e:
                print(f"Error comprimiendo la foto {foto.NombreArchivo}: {e}")
                # Si falla, devolvemos la original para no romper el flujo
                pass 
                
        fotos_optimizadas.append(foto)
        
    return fotos_optimizadas

# ==========================================
# REPARACIÓN DE TAGS ROTOS (Con el Escudo)
# ==========================================
def reparar_tags_rotos(documento_docx):
    def fusionar_parrafo(p):
        texto_completo = "".join(run.text for run in p.runs)
        
        # EL ESCUDO: Solo entramos si el párrafo contiene "{%" y "tr"
        # Esto deja tranquilas a variables normales como {{ Metadatos.Fig }}
        if "{%" in texto_completo and "tr" in texto_completo:
            
            # TU SOLUCIÓN: Quitamos el espacio y forzamos el {%tr junto
            texto_corregido = texto_completo.replace("{% tr", "{%tr").replace("{%  tr", "{%tr")
            
            if len(p.runs) > 0:
                p.runs[0].text = texto_corregido
                for run in p.runs[1:]:
                    run.text = ""

    for p in documento_docx.paragraphs:
        fusionar_parrafo(p)
    for tabla in documento_docx.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    fusionar_parrafo(p)

# ==========================================
# MAGIA: ESCANER DE IMÁGENES Y LIMPIEZA DE NULOS
# ==========================================
def procesar_datos_rec(sub_contexto, doc):
    if isinstance(sub_contexto, dict):
        for k, v in list(sub_contexto.items()):
            # 1. Limpiar Nulos (si llega null, se convierte en texto vacío)
            if v is None:
                sub_contexto[k] = ""
            
            # 2. Procesar Imágenes
            elif isinstance(v, str) and ("base64," in v or v.startswith("data:image")):
                try:
                    foto_str = v.split("base64,")[-1]
                    img_bytes_raw = base64.b64decode(foto_str)

                    pil_img = Image.open(BytesIO(img_bytes_raw)).convert("RGB")
                    buffer_corregido = BytesIO()
                    pil_img.save(buffer_corregido, format="JPEG", dpi=(96, 96))
                    buffer_corregido.seek(0)

                    sub_contexto[k] = InlineImage(doc, buffer_corregido, width=Mm(50))
                except Exception as e:
                    print(f"No se pudo procesar la imagen en '{k}': {e}")
                    sub_contexto[k] = ""
            else:
                procesar_datos_rec(v, doc)
                
    elif isinstance(sub_contexto, list):
        for i in range(len(sub_contexto)):
            if sub_contexto[i] is None:
                sub_contexto[i] = ""
            else:
                procesar_datos_rec(sub_contexto[i], doc)

# ==========================================
# ENDPOINT UNIVERSAL
# ==========================================
@app.post("/api/v1/generar-reporte/")
def generar_reporte_universal(payload: PayloadUniversal):
    try:
        ruta_plantilla = f"templates/{payload.template_name}.docx"
        
        try:
            doc_puro = Document(ruta_plantilla)
        except Exception:
            raise HTTPException(status_code=404, detail=f"La plantilla '{payload.template_name}.docx' no existe.")

        reparar_tags_rotos(doc_puro)

        buffer_reparado = BytesIO()
        doc_puro.save(buffer_reparado)
        buffer_reparado.seek(0) 

        doc = DocxTemplate(buffer_reparado)

        contexto = payload.data
        procesar_datos_rec(contexto, doc)

        doc.render(contexto)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        nombre_descarga = f"Reporte_{payload.template_name}.docx"
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre_descarga}"'}
        )

    except Exception as e:
        print(f"Error crítico en el servidor: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# ENDPOINT PARA MANTENER DESPIERTO (PING)
# ==========================================
@app.get("/api/v1/ping/")
def ping_server():
    return {"status": "ok", "message": "Servidor de reportes despierto y listo."}